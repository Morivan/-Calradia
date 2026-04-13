from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_PATH = Path(r"E:\Project\output\doc\Раздел_тестирование_Кузница_Кальрадия.docx")


TEST_CASES = [
    ("1", "Загрузка агрегированных данных системы", "Выполнить запрос на загрузку стартовых данных приложения", "Система возвращает список товаров, клиентов CRM, материалов, финансовых расчётов, маркетинговых публикаций и внешних ссылок"),
    ("2", "Создание отзыва к товару", "Отправить данные отзыва: автор, текст, оценка, дата", "Отзыв сохраняется в системе и становится доступным в карточке товара"),
    ("3", "Загрузка каталога продукции", "Открыть модуль каталога", "На экране отображается список товаров с изображением, названием, материалом и ценой"),
    ("4", "Открытие карточки товара", "Выбрать товар в каталоге и перейти к просмотру деталей", "Открывается детальная карточка с описанием, характеристиками, галереей и кнопкой связи"),
    ("5", "Синхронизация справочника материалов из Google Sheets", "Запустить синхронизацию финансового модуля", "В систему загружается справочник материалов и история расчётов"),
    ("6", "Обработка ошибки синхронизации Google Sheets", "Выполнить синхронизацию при недоступной таблице", "Система возвращает корректное сообщение об ошибке без аварийного завершения"),
    ("7", "Ответ клиенту через CRM", "Отправить сообщение клиенту из интерфейса CRM", "Сообщение сохраняется в системе и, для Telegram-клиента, передаётся через Telegram API"),
    ("8", "Создание заказа из CRM", "Передать данные клиента, товара, состава заказа и стоимости", "В системе создаётся заказ с заданным статусом"),
    ("9", "Ручная публикация маркетингового поста", "Выбрать публикацию и запустить публикацию в Telegram", "Публикация отправляется в Telegram, статус меняется на «Опубликован»"),
    ("10", "Автоматическая публикация по расписанию", "Запустить команду публикации запланированных постов", "Все публикации, время которых наступило, публикуются автоматически"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)


def add_paragraph(document: Document, text: str, bold: bool = False, first_line_indent_cm: float = 1.25) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14 if level == 1 else 13)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.left_indent = Cm(1.25)
        paragraph.paragraph_format.first_line_indent = Cm(-0.5)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(f"- {item}")
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)


def add_test_cases_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["№", "Наименование тест-кейса", "Входные данные / действия", "Ожидаемый результат"]
    widths = [Cm(1.2), Cm(5.4), Cm(5.8), Cm(5.8)]

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        header_cells[index].width = widths[index]
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(header_cells[index], "D9EAD3")
        paragraph = header_cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)

    for row in TEST_CASES:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            cells[index].width = widths[index]
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(11)


def build_document() -> Document:
    document = Document()
    style_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run("Раздел дипломной работы\n«Тестирование программного комплекса “Кузница Кальрадия”»")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title_run.font.size = Pt(14)

    add_heading(document, "5 Тестирование программного комплекса")
    add_paragraph(
        document,
        "Целью тестирования разработанного программного комплекса «Кузница Кальрадия» являлась проверка корректности реализации основных функциональных сценариев, устойчивости работы программных модулей, а также оценка производительности системы при типовой нагрузке. Тестирование проводилось для клиентской и серверной частей приложения с использованием автоматизированных и ручных методов проверки.",
    )
    add_paragraph(document, "В рамках тестирования решались следующие задачи:")
    add_bullets(
        document,
        [
            "проверка корректности работы ключевых бизнес-сценариев системы;",
            "проверка отдельных программных компонентов на уровне модулей;",
            "проверка взаимодействия между клиентской и серверной частями;",
            "оценка устойчивости API при многократных параллельных запросах;",
            "подтверждение готовности функционального прототипа к демонстрации заказчику.",
        ],
    )

    add_heading(document, "5.1 Виды проведённого тестирования", level=2)
    add_paragraph(
        document,
        "Для оценки качества разработанной системы были использованы следующие виды тестирования: модульное тестирование, интеграционное и сквозное тестирование, нагрузочное тестирование, а также тестирование сборки клиентской части.",
    )
    add_paragraph(
        document,
        "Модульное тестирование применялось для проверки отдельных функций и сервисов системы без учёта полного пользовательского сценария. В ходе данного этапа проверялись функции обработки ссылок Google Sheets, логика Telegram-интеграции, функции маркетингового модуля, а также корректность сериализации данных API.",
    )
    add_paragraph(
        document,
        "Интеграционное и сквозное тестирование использовалось для проверки взаимодействия между основными модулями системы и серверными интерфейсами. Данный этап позволил подтвердить корректность выполнения пользовательских сценариев, включающих каталог продукции, CRM, финансовый модуль и модуль маркетинга.",
    )
    add_paragraph(
        document,
        "Нагрузочное тестирование применялось для оценки производительности ключевого API-метода, обеспечивающего загрузку агрегированных данных системы. Проверялась способность серверной части обрабатывать серию параллельных запросов без ошибок и существенного ухудшения времени отклика.",
    )
    add_paragraph(
        document,
        "Тестирование сборки клиентской части проводилось для подтверждения корректной компиляции frontend-приложения и готовности системы к развёртыванию.",
    )

    add_heading(document, "5.2 Таблица тест-кейсов для приложения", level=2)
    add_test_cases_table(document)

    add_heading(document, "5.3 Описание тестовой среды", level=2)
    add_paragraph(
        document,
        "Тестирование программного комплекса «Кузница Кальрадия» проводилось в локальной среде разработки с использованием клиентской и серверной частей приложения. В качестве серверной платформы использовался интерпретатор Python 3.13 и фреймворк Django. Клиентская часть приложения разрабатывалась и собиралась с использованием React, TypeScript и Vite.",
    )
    add_paragraph(document, "В ходе тестирования применялись следующие программные средства:")
    add_bullets(
        document,
        [
            "Python 3.13;",
            "Django и Django REST Framework;",
            "SQLite в качестве тестовой базы данных;",
            "Node.js и npm для сборки клиентской части;",
            "Vite для разработки и production-сборки frontend-приложения;",
            "встроенный тестовый фреймворк Django;",
            "PowerShell в качестве командной оболочки;",
            "веб-браузер Google Chrome для ручной проверки интерфейсов.",
        ],
    )
    add_paragraph(
        document,
        "Для интеграционных проверок использовались Telegram Bot API для обмена сообщениями и публикации маркетинговых сообщений, а также Google Sheets в качестве внешнего источника данных финансового модуля.",
    )
    add_paragraph(
        document,
        "Нагрузочное тестирование проводилось для серверного интерфейса /api/bootstrap/ с использованием отдельного Python-скрипта, формирующего серию параллельных HTTP-запросов к локально запущенному backend-приложению.",
    )
    add_paragraph(
        document,
        "Тестирование выполнялось на персональном компьютере под управлением операционной системы Windows. Проверка включала автоматизированные тесты серверной части, сборку клиентского приложения, а также ручной контроль пользовательских сценариев в браузере.",
    )

    add_heading(document, "5.4 Результаты тестирования", level=2)
    add_paragraph(
        document,
        "Автоматизированное тестирование серверной части выполнялось с использованием встроенных средств тестирования Django. В результате выполнения тестового набора установлено, что все предусмотренные проверки завершились успешно.",
    )
    add_bullets(
        document,
        [
            "модульные тесты: 20 из 20 успешно;",
            "сквозные тест-кейсы: 10 из 10 успешно;",
            "общее количество успешно выполненных автоматизированных проверок: 30 из 30;",
            "production-сборка клиентской части выполнена успешно.",
        ],
    )
    add_paragraph(
        document,
        "Нагрузочное тестирование проводилось для серверного интерфейса /api/bootstrap/. В рамках теста было выполнено 200 запросов при уровне параллелизма 20 соединений.",
    )
    add_bullets(
        document,
        [
            "общее число запросов: 200;",
            "успешных ответов: 200;",
            "количество ошибок: 0;",
            "минимальное время ответа: 92,45 мс;",
            "среднее время ответа: 255,63 мс;",
            "95-й перцентиль времени ответа: 307,38 мс;",
            "максимальное время ответа: 340,18 мс;",
            "пропускная способность: 76,74 запросов в секунду.",
        ],
    )

    add_heading(document, "5.5 Вывод по результатам тестирования", level=2)
    add_paragraph(
        document,
        "По результатам проведённого тестирования можно сделать вывод, что разработанный программный комплекс «Кузница Кальрадия» корректно реализует основные функциональные сценарии, предусмотренные техническим заданием. Система успешно проходит модульные, сквозные и нагрузочные проверки, демонстрирует стабильную работу основных интерфейсов и готова к использованию в качестве функционального прототипа и дальнейшему развитию.",
    )

    return document


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
